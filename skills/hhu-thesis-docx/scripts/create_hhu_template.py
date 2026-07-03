#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成河海大学学位论文 Word 基础骨架（样式 + 分节 + 页眉页脚）。
用法:
    python create_hhu_template.py [输出路径] [--degree doctor|academicmaster|professionalmaster|bachelor]

生成的文档包含 3 个节:
  节1: 前置部分(封面等) — 无页眉页脚无页码
  节2: 摘要/目录等     — 罗马页码 I, II...
  节3: 正文            — 阿拉伯页码从 1 起
调用方在对应节内填充内容即可。也可 import 本模块复用工具函数。
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_TEXT = {
    'doctor': '河海大学博士学位论文',
    'academicmaster': '河海大学硕士学位论文',
    'professionalmaster': '河海大学硕士学位论文',
    'bachelor': '河海大学学士学位论文',
}

# ---------- 基础工具 ----------

def set_run_font(run, zh='宋体', en='Times New Roman', size=12, bold=False):
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), zh)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_style_font(style, zh, en, size, bold):
    style.font.name = en
    style.element.get_or_add_rPr()
    rFonts = style.element.rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); style.element.rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), zh)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)

def set_exact_spacing(pf, pts=20, before=0, after=0):
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pts)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def set_first_line_chars(paragraph, chars=200):
    """首行缩进按字符数: 200 = 2字符; 0 = 不缩进"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))
    ind.set(qn('w:firstLine'), str(int(chars * 2.4)))  # 兜底 twips(12pt字)

def style_first_line_chars(style, chars=200):
    pPr = style.element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars))

def add_field(paragraph, instr, font_size=9):
    """插入域代码（PAGE / STYLEREF / TOC 等）"""
    run = paragraph.add_run()
    fld_b = OxmlElement('w:fldChar'); fld_b.set(qn('w:fldCharType'), 'begin')
    instr_el = OxmlElement('w:instrText'); instr_el.set(qn('xml:space'), 'preserve')
    instr_el.text = instr
    fld_e = OxmlElement('w:fldChar'); fld_e.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_b); run._r.append(instr_el); run._r.append(fld_e)
    set_run_font(run, '宋体', 'Times New Roman', font_size)
    return run

def set_header_double_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'thickThinSmallGap')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom); pPr.append(pBdr)

def set_page_number_format(section, fmt, start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType'); sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

def setup_page(section):
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.5)
    section.left_margin = section.right_margin = Cm(2.7)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.8)

# ---------- 样式定义 ----------

def define_styles(doc):
    # 正文
    st = doc.styles['Normal']
    set_style_font(st, '宋体', 'Times New Roman', 12, False)
    set_exact_spacing(st.paragraph_format, 20)
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    try:
        body = doc.styles.add_style('HHU正文', 1)  # paragraph style
    except Exception:
        body = doc.styles['HHU正文']
    body.base_style = doc.styles['Normal']
    set_style_font(body, '宋体', 'Times New Roman', 12, False)
    set_exact_spacing(body.paragraph_format, 20)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_first_line_chars(body, 200)

    # 章: 黑体三号加粗居中, 段前16磅段后32磅
    h1 = doc.styles['Heading 1']
    set_style_font(h1, '黑体', 'Times New Roman', 16, True)
    set_exact_spacing(h1.paragraph_format, 20, 16, 32)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.page_break_before = True
    style_first_line_chars(h1, 0)

    # 节: 黑体四号加粗左对齐
    h2 = doc.styles['Heading 2']
    set_style_font(h2, '黑体', 'Times New Roman', 14, True)
    set_exact_spacing(h2.paragraph_format, 20, 10, 10)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_first_line_chars(h2, 0)

    # 条/款: 黑体小四加粗左对齐
    for name in ('Heading 3', 'Heading 4'):
        h = doc.styles[name]
        set_style_font(h, '黑体', 'Times New Roman', 12, True)
        set_exact_spacing(h.paragraph_format, 20, 10, 10)
        h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style_first_line_chars(h, 0)

    # 图题/表题: 五号加粗居中
    for nm in ('HHU图题', 'HHU表题'):
        try:
            cap = doc.styles.add_style(nm, 1)
        except Exception:
            cap = doc.styles[nm]
        set_style_font(cap, '宋体', 'Times New Roman', 10.5, True)
        set_exact_spacing(cap.paragraph_format, 20, 6, 6)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_first_line_chars(cap, 0)

    # 参考文献条目: 五号, 悬挂缩进
    try:
        ref = doc.styles.add_style('HHU参考文献', 1)
    except Exception:
        ref = doc.styles['HHU参考文献']
    set_style_font(ref, '宋体', 'Times New Roman', 10.5, False)
    set_exact_spacing(ref.paragraph_format, 16)
    ref.paragraph_format.left_indent = Cm(0.74)
    ref.paragraph_format.first_line_indent = Cm(-0.74)  # 悬挂
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ---------- 页眉页脚 ----------

def enable_even_odd_headers(doc):
    settings = doc.settings.element
    if settings.find(qn('w:evenAndOddHeaders')) is None:
        settings.append(OxmlElement('w:evenAndOddHeaders'))

def build_headers_footers(section, header_even_text, odd_use_styleref=True,
                          odd_text=None, heading_style_name='Heading 1'):
    """配置一个节的奇偶页眉 + 居中页码页脚"""
    for h in (section.header, section.even_page_header,
              section.footer, section.even_page_footer):
        h.is_linked_to_previous = False
        for p in list(h.paragraphs[1:]):
            p._p.getparent().remove(p._p)
        h.paragraphs[0].text = ''

    # 偶数页页眉: 固定校名文字
    pe = section.even_page_header.paragraphs[0]
    pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pe.add_run(header_even_text); set_run_font(r, '宋体', 'Times New Roman', 9)
    set_header_double_border(pe)
    set_exact_spacing(pe.paragraph_format, 14)

    # 奇数页页眉: 当前章标题(STYLEREF域)或固定文字
    po = section.header.paragraphs[0]
    po.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if odd_use_styleref:
        add_field(po, f' STYLEREF "{heading_style_name}" \\* MERGEFORMAT ', 9)
    elif odd_text:
        r = po.add_run(odd_text); set_run_font(r, '宋体', 'Times New Roman', 9)
    set_header_double_border(po)
    set_exact_spacing(po.paragraph_format, 14)

    # 页脚: 居中页码 宋体小五
    for f in (section.footer, section.even_page_footer):
        p = f.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(p, ' PAGE \\* MERGEFORMAT ', 9)

def clear_headers_footers(section):
    for h in (section.header, section.even_page_header,
              section.footer, section.even_page_footer):
        h.is_linked_to_previous = False
        for p in list(h.paragraphs[1:]):
            p._p.getparent().remove(p._p)
        h.paragraphs[0].text = ''

# ---------- 主流程 ----------

def create_base_document(degree='doctor'):
    doc = Document()
    define_styles(doc)
    enable_even_odd_headers(doc)
    header_text = HEADER_TEXT[degree]

    # 节1: 封面等前置页(无页眉页脚)
    sec1 = doc.sections[0]
    setup_page(sec1)
    sec1.different_first_page_header_footer = False
    clear_headers_footers(sec1)
    p = doc.add_paragraph(); r = p.add_run('【节1：封面/题名页/英文扉页/声明页 — 在此填充，每页之间用分页符】')
    set_run_font(r, '宋体', 'Times New Roman', 12)

    # 节2: 摘要/目录等, 罗马页码
    sec2 = doc.add_section(WD_SECTION_START.ODD_PAGE)
    setup_page(sec2)
    set_page_number_format(sec2, 'upperRoman', 1)
    build_headers_footers(sec2, header_text)
    p = doc.add_paragraph(); r = p.add_run('【节2：前言/摘要/Abstract/符号表/目录/插图清单/附表清单 — 罗马页码】')
    set_run_font(r, '宋体', 'Times New Roman', 12)

    # 节3: 正文, 阿拉伯页码从1
    sec3 = doc.add_section(WD_SECTION_START.ODD_PAGE)
    setup_page(sec3)
    set_page_number_format(sec3, 'decimal', 1)
    build_headers_footers(sec3, header_text)
    p = doc.add_paragraph(); r = p.add_run('【节3：正文各章/参考文献/致谢/附录 — 阿拉伯页码从1】')
    set_run_font(r, '宋体', 'Times New Roman', 12)

    return doc

if __name__ == '__main__':
    out = sys.argv[1] if len(sys.argv) > 1 and not sys.argv[1].startswith('--') else 'hhu_thesis_base.docx'
    degree = 'doctor'
    if '--degree' in sys.argv:
        degree = sys.argv[sys.argv.index('--degree') + 1]
    doc = create_base_document(degree)
    doc.save(out)
    print(f'已生成: {out} (学位类型: {degree})')
