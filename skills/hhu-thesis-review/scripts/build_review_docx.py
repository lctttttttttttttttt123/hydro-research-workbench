#!/usr/bin/env python3
"""
把审稿结果生成为一个清晰可读的 Word（.docx）文档。

用法:
    python build_review_docx.py <spec.json> <输出路径.docx>

spec.json 结构见 references/docx-output.md。核心字段:
{
  "title": "审稿意见：摘要",                 # 文档标题
  "scope": "本次审查范围：摘要",             # 一句话说明审了哪几块（可选）
  "overview": "一段简短总览……",            # 可选，放文档开头
  "cross_section_issues": [                  # 总体/跨段问题区（可空）
     {"severity": "硬伤", "text": "背景段在前言/摘要/结论三处逐句复用……"}
  ],
  "sentence_edits": [                        # 逐句修改表（可空）
     {"locate": "揭示了闸泵调度对流场的影响规律。",
      "issue": "【硬伤】实测类结论无任何数值，读者无法判断结论。",
      "revised": "闸泵联合调度使主干河道平均流速提高约 X%。",
      "reason": "把定性动词换成可被数据支撑的判断句，并补关键数字。"}
  ],
  "single_block_footnote": "摘要开头属典型背景模板，请自查前言/结论是否照搬同段。"  # 可选
}

设计要点（已踩坑，勿从零写）:
- 中文必须设 eastAsia 字体，否则 Word 里中文回退。
- 表格四周加边框（sz 单位为 1/8 磅）。
- 文档不追求论文正文格式，清晰可读即可，但中文字体务必正确。
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ZH = "宋体"
ZH_BOLD = "黑体"
EN = "Times New Roman"

SEV_COLOR = {
    "硬伤": RGBColor(0xC0, 0x00, 0x00),   # 深红
    "建议": RGBColor(0xC0, 0x55, 0x00),   # 橙
    "可选": RGBColor(0x55, 0x55, 0x55),   # 灰
}


def set_run_font(run, zh=ZH, en=EN, size_pt=12, bold=False, color=None):
    """设置中英文字体（中文必须走 eastAsia，否则回退）。"""
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), zh)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", size=12, bold=False, zh=ZH, align=None,
             space_before=0, space_after=6, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(20)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, zh=zh, size_pt=size, bold=bold, color=color)
    return p


def add_severity_run(paragraph, severity, size=12):
    """在段首加一个带色的【严重度】标签 run。"""
    tag = paragraph.add_run(f"【{severity}】")
    set_run_font(tag, zh=ZH_BOLD, size_pt=size, bold=True,
                 color=SEV_COLOR.get(severity, RGBColor(0, 0, 0)))


def set_cell_text(cell, text, size=10.5, bold=False, zh=ZH, color=None,
                  severity_prefix=None):
    """写单元格文本，正确设置中文字体。可选在前面加带色严重度标签。"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    if severity_prefix:
        add_severity_run(p, severity_prefix, size=size)
    r = p.add_run(text)
    set_run_font(r, zh=zh, size_pt=size, bold=bold, color=color)


def _set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")          # 0.75 磅 (sz=1/8磅)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)


def shade_cell(cell, fill="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def build(spec, out_path):
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = EN
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), ZH)

    add_para(doc, spec.get("title", "审稿意见"), size=18, bold=True, zh=ZH_BOLD,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    if spec.get("scope"):
        add_para(doc, spec["scope"], size=10.5, zh=ZH,
                 align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55),
                 space_after=4)
    add_para(doc,
             "说明：本文档中“修改后”列及示范句里的数字均为示意占位（如 X%、Δ天、相关系数 X），请替换为你的实测值后再使用。",
             size=9, zh=ZH, color=RGBColor(0xC0, 0x55, 0x00), space_after=10)

    if spec.get("overview"):
        add_para(doc, "总览", size=13, bold=True, zh=ZH_BOLD, space_before=4, space_after=4)
        add_para(doc, spec["overview"], size=11, zh=ZH, space_after=10)

    add_para(doc, "一、总体与跨段问题", size=14, bold=True, zh=ZH_BOLD,
             space_before=6, space_after=4)
    issues = spec.get("cross_section_issues") or []
    if not issues:
        add_para(doc, "（本次未发现总体/跨段结构性问题。）", size=11, zh=ZH,
                 color=RGBColor(0x55, 0x55, 0x55), space_after=8)
    else:
        for it in issues:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.line_spacing = Pt(20)
            p.paragraph_format.space_after = Pt(4)
            add_severity_run(p, it.get("severity", "建议"), size=11)
            r = p.add_run(it.get("text", ""))
            set_run_font(r, zh=ZH, size_pt=11)

    add_para(doc, "二、逐句修改表", size=14, bold=True, zh=ZH_BOLD,
             space_before=10, space_after=2)
    add_para(doc, "（只列有问题的句子；没有问题的句子不在表内。）", size=9, zh=ZH,
             color=RGBColor(0x55, 0x55, 0x55), space_after=6)

    edits = spec.get("sentence_edits") or []
    if not edits:
        add_para(doc, "（本次未发现需逐句修改的问题。）", size=11, zh=ZH,
                 color=RGBColor(0x55, 0x55, 0x55), space_after=8)
    else:
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        widths = [Cm(4.2), Cm(4.6), Cm(4.6), Cm(4.2)]
        headers = ["原文定位", "问题", "修改后", "修改理由"]
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, size=10.5, bold=True, zh=ZH_BOLD)
            shade_cell(hdr[i])
            _set_cell_border(hdr[i])
            hdr[i].width = widths[i]
        for e in edits:
            row = table.add_row().cells
            set_cell_text(row[0], e.get("locate", ""), size=10.5, zh=ZH)
            issue = e.get("issue", "")
            sev = None
            for s in ("硬伤", "建议", "可选"):
                if issue.startswith(f"【{s}】"):
                    sev = s
                    issue = issue[len(f"【{s}】"):]
                    break
            set_cell_text(row[1], issue, size=10.5, zh=ZH, severity_prefix=sev)
            set_cell_text(row[2], e.get("revised", ""), size=10.5, zh=ZH)
            set_cell_text(row[3], e.get("reason", ""), size=10.5, zh=ZH)
            for i in range(4):
                _set_cell_border(row[i])
                row[i].width = widths[i]

    if spec.get("single_block_footnote"):
        add_para(doc, "", space_after=2)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(18)
        r0 = p.add_run("脚注：")
        set_run_font(r0, zh=ZH_BOLD, size_pt=9, bold=True, color=RGBColor(0x55, 0x55, 0x55))
        r1 = p.add_run(spec["single_block_footnote"])
        set_run_font(r1, zh=ZH, size_pt=9, color=RGBColor(0x55, 0x55, 0x55))

    doc.save(out_path)
    print(f"已生成: {out_path}")


def main():
    if len(sys.argv) != 3:
        print("用法: python build_review_docx.py <spec.json> <输出.docx>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        spec = json.load(f)
    build(spec, sys.argv[2])


if __name__ == "__main__":
    main()
